"""
Build a real .docx resume from structured content using python-docx.

The previous Word "export" wrapped the template HTML string in a Blob with a
.docx MIME type. That is not a Word document: Word shows a
"format and extension don't match" warning, links are unreliable, and there is
no real pagination. This builds a genuine OOXML document with headings,
real hyperlinks, and clean single-column layout that any ATS can parse.

Input shape mirrors the frontend `TemplateInput`:
    {
      "candidate_info": {name, email, phone, linkedin?, portfolio?},
      "resume": {
        "summary": str,
        "experiences": [{role, company, location?, startDate, endDate?, bullets[]}],
        "projects":    [{title, link?: {url,label?} | str, bullets[]}],
        "education":   [{school, degree, field, location?, endDate}],
        "skills":      [{category, skills[]}]
      }
    }
"""
from __future__ import annotations

import io
from typing import Any, Optional

from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

NAVY = RGBColor(0x1F, 0x3A, 0x5F)
LINK_BLUE = RGBColor(0x1A, 0x4D, 0xA1)
GREY = RGBColor(0x44, 0x44, 0x44)

_MONTHS = [
    "January", "February", "March", "April", "May", "June",
    "July", "August", "September", "October", "November", "December",
]


def _fmt_date(value: Optional[str]) -> str:
    if not value:
        return "Present"
    v = str(value).strip()
    if len(v) == 4 and v.isdigit():
        return v
    if len(v) == 7 and v[4] == "-":
        try:
            y, m = v.split("-")
            idx = int(m) - 1
            if 0 <= idx < 12:
                return f"{_MONTHS[idx]} {y}"
        except Exception:
            pass
    return v


def _date_range(start: Optional[str], end: Optional[str]) -> str:
    return f"{_fmt_date(start)} - {_fmt_date(end)}"


def _normalize_link(link: Any) -> Optional[dict]:
    """Accept either a bare URL string or a {url, label} object."""
    if not link:
        return None
    if isinstance(link, str):
        return {"url": link, "label": link}
    if isinstance(link, dict) and link.get("url"):
        return {"url": link["url"], "label": link.get("label") or link["url"]}
    return None


def _ensure_url(url: str) -> str:
    if url.startswith(("http://", "https://", "mailto:")):
        return url
    if "@" in url and "/" not in url:
        return f"mailto:{url}"
    return f"https://{url}"


def _add_hyperlink(paragraph, url: str, text: str, color: RGBColor = LINK_BLUE):
    """Add a real clickable hyperlink run to a paragraph (python-docx has no
    high-level API for this)."""
    part = paragraph.part
    r_id = part.relate_to(
        _ensure_url(url),
        "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink",
        is_external=True,
    )
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), r_id)

    new_run = OxmlElement("w:r")
    rpr = OxmlElement("w:rPr")

    c = OxmlElement("w:color")
    c.set(qn("w:val"), "%02X%02X%02X" % (color[0], color[1], color[2]))
    rpr.append(c)
    u = OxmlElement("w:u")
    u.set(qn("w:val"), "single")
    rpr.append(u)

    new_run.append(rpr)
    t = OxmlElement("w:t")
    t.text = text
    new_run.append(t)
    hyperlink.append(new_run)
    paragraph._p.append(hyperlink)
    return hyperlink


def _section_heading(doc: Document, title: str):
    p = doc.add_paragraph()
    p.space_before = Pt(8)
    run = p.add_run(title.upper())
    run.bold = True
    run.font.size = Pt(11)
    run.font.color.rgb = NAVY
    ppr = p._p.get_or_add_pPr()
    borders = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "6")
    bottom.set(qn("w:space"), "2")
    bottom.set(qn("w:color"), "9AA3AD")
    borders.append(bottom)
    ppr.append(borders)
    return p


def build_resume_docx(data: dict) -> bytes:
    ci = (data or {}).get("candidate_info", {}) or {}
    r = (data or {}).get("resume", {}) or {}

    doc = Document()
    for section in doc.sections:
        section.top_margin = Pt(40)
        section.bottom_margin = Pt(40)
        section.left_margin = Pt(54)
        section.right_margin = Pt(54)

    # ---- Header ----
    name_p = doc.add_paragraph()
    name_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    name_run = name_p.add_run(ci.get("name") or "Resume")
    name_run.bold = True
    name_run.font.size = Pt(20)
    name_run.font.color.rgb = NAVY

    contact_p = doc.add_paragraph()
    contact_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    first = True

    def _sep():
        nonlocal first
        if not first:
            contact_p.add_run("  |  ").font.color.rgb = GREY
        first = False

    if ci.get("email"):
        _sep()
        _add_hyperlink(contact_p, ci["email"], ci["email"], GREY)
    if ci.get("phone"):
        _sep()
        contact_p.add_run(ci["phone"]).font.color.rgb = GREY
    if ci.get("linkedin"):
        _sep()
        _add_hyperlink(contact_p, ci["linkedin"], ci["linkedin"], GREY)
    if ci.get("portfolio"):
        _sep()
        _add_hyperlink(contact_p, ci["portfolio"], ci["portfolio"], GREY)

    # ---- Summary ----
    if r.get("summary"):
        _section_heading(doc, "Summary")
        doc.add_paragraph(r["summary"])

    # ---- Experience ----
    experiences = r.get("experiences") or []
    if experiences:
        _section_heading(doc, "Experience")
        for exp in experiences:
            role_p = doc.add_paragraph()
            role_run = role_p.add_run(exp.get("role") or "")
            role_run.bold = True
            role_run.font.size = Pt(10.5)

            meta_p = doc.add_paragraph()
            company = exp.get("company") or ""
            loc = exp.get("location")
            date_txt = _date_range(exp.get("startDate"), exp.get("endDate"))
            if loc:
                date_txt = f"{date_txt}, {loc}"
            cr = meta_p.add_run(company)
            cr.bold = True
            meta_p.add_run("    " + date_txt).italic = True

            for b in exp.get("bullets") or []:
                doc.add_paragraph(str(b), style="List Bullet")

    # ---- Projects ----
    projects = r.get("projects") or []
    if projects:
        _section_heading(doc, "Projects")
        for proj in projects:
            title_p = doc.add_paragraph()
            title_p.add_run(proj.get("title") or "Project").bold = True
            link = _normalize_link(proj.get("link"))
            if link:
                title_p.add_run("  ")
                _add_hyperlink(title_p, link["url"], link["label"])
            for b in proj.get("bullets") or []:
                doc.add_paragraph(str(b), style="List Bullet")

    # ---- Education ----
    education = r.get("education") or []
    if education:
        _section_heading(doc, "Education")
        for edu in education:
            degree = (edu.get("degree") or "").strip()
            field = (edu.get("field") or "").strip()
            if degree and field and field.lower() not in degree.lower():
                full_degree = f"{degree} in {field}"
            else:
                full_degree = degree or field
            dp = doc.add_paragraph()
            dp.add_run(full_degree).bold = True
            meta = " • ".join(
                x for x in [edu.get("school"), edu.get("location"), _fmt_date(edu.get("endDate")) if edu.get("endDate") else None] if x
            )
            if meta:
                doc.add_paragraph(meta)

    # ---- Skills ----
    skills = r.get("skills") or []
    if skills:
        _section_heading(doc, "Skills")
        for cat in skills:
            sp = doc.add_paragraph()
            label = cat.get("category")
            if label:
                sp.add_run(f"{label}: ").bold = True
            sp.add_run(", ".join(cat.get("skills") or []))

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()
