"""
Automated smoke + stress tests for the resume export pipeline.

Requirement under test:
  * PDF and DOCX downloads are ACCURATE (real text, real links, valid files).
  * When data spans MULTIPLE pages, indentation / padding / spacing must stay
    consistent on every page (no flush-to-edge, no sliced content).

Run:  python test_export_smoke.py
(Render the fixtures first:  cd ../resume-ai-frontend && node scripts/_smoke-render.mjs)

Exit code 0 = all assertions passed.
"""
import io
import os
import re
import sys
import glob
import zipfile

sys.path.insert(0, os.path.dirname(__file__))
from app.utils.pdf_renderer import render_html_to_pdf            # noqa: E402
from app.utils.docx_builder import build_resume_docx            # noqa: E402

import fitz                                                      # noqa: E402
from pypdf import PdfReader                                      # noqa: E402
from docx import Document                                        # noqa: E402

FIX = os.path.abspath(os.path.join(
    os.path.dirname(__file__), "..", "resume-ai-frontend", "scripts", "output", "_smoke"))

# margin tolerances (mm)
MIN_TOP_MARGIN_MM = 8.0      # page 2+ must NOT be flush to the paper edge
MAX_MARGIN_SPREAD_MM = 6.0   # top/left margins must be consistent across pages
DPI = 150
PX_PER_MM = DPI / 25.4

PASS, FAIL = [], []


def check(name, cond, detail=""):
    (PASS if cond else FAIL).append(name)
    print(f"  {'PASS' if cond else 'FAIL'}  {name}{('  — ' + detail) if detail and not cond else ''}")


def _dark_bleed_band_fraction(page):
    """Fraction of rows in the lower part of the page that are a FULL-WIDTH dark
    band — i.e. a solid colored block bleeding into the empty area below the
    content (a template's dark <html> background showing through).

    Keyed on full-width darkness so a legitimate narrow left sidebar column
    (only ~30% of the width) does NOT count, but an edge-to-edge dark fill does.
    """
    pix = page.get_pixmap(dpi=100)
    w, h, n, s, stride = pix.width, pix.height, pix.n, pix.samples, pix.stride
    y0 = int(h * 0.65)               # examine the bottom ~35%
    band_rows = bleed_rows = 0
    for y in range(y0, h):
        base = y * stride
        dark = wide = 0
        for x in range(0, w, 2):
            idx = base + x * n
            lum = 0.299 * s[idx] + 0.587 * s[idx + 1] + 0.114 * s[idx + 2]
            wide += 1
            if lum < 90:
                dark += 1
        band_rows += 1
        if dark / wide > 0.85:       # this row is dark across (almost) the full width
            bleed_rows += 1
    return (bleed_rows / band_rows) if band_rows else 0.0


def _ink_bounds(page):
    """Return (top_mm, left_mm) of the first inked pixel on the page."""
    pix = page.get_pixmap(dpi=DPI)
    w, h, n, s, stride = pix.width, pix.height, pix.n, pix.samples, pix.stride
    top = None
    left = w
    for y in range(h):
        base = y * stride
        row_ink_left = None
        for x in range(w):
            idx = base + x * n
            if s[idx] < 235 or s[idx + 1] < 235 or s[idx + 2] < 235:
                row_ink_left = x
                break
        if row_ink_left is not None:
            if top is None:
                top = y
            left = min(left, row_ink_left)
    return (top / PX_PER_MM if top is not None else 999.0,
            left / PX_PER_MM if left < w else 999.0)


def extract_text_links(pdf_bytes):
    r = PdfReader(io.BytesIO(pdf_bytes))
    text = "\n".join((p.extract_text() or "") for p in r.pages)
    links = []
    for p in r.pages:
        for a in (p.get("/Annots") or []):
            uri = a.get_object().get("/A", {}).get("/URI")
            if uri:
                links.append(uri)
    return len(r.pages), text, links


def test_pdf(slug, profile, html, multipage):
    label = f"[{slug}/{profile}] PDF"
    try:
        pdf = render_html_to_pdf(html)
    except Exception as e:
        check(f"{label} renders", False, str(e)[:120])
        return
    check(f"{label} is a valid PDF", pdf[:5] == b"%PDF-")
    pages, text, links = extract_text_links(pdf)
    check(f"{label} has selectable text", len(text.strip()) > 400, f"len={len(text.strip())}")
    check(f"{label} has clickable link(s)", len(links) >= 1, f"links={len(links)}")
    # date not garbled (the html2canvas 'N - vember' artifact must be gone)
    garbled = bool(re.search(r"N\s*[-–]\s*vember", text))
    check(f"{label} date not garbled", not garbled)

    doc = fitz.open(stream=pdf, filetype="pdf")
    # No template's background may bleed a full-width colored block into the
    # empty lower area of any page (the creative-bold / left-sidebar dark <html>
    # bg bug). 5% tolerance for tiny baseline rules.
    worst = max(_dark_bleed_band_fraction(pg) for pg in doc)
    check(f"{label} no colored-fill bleed in empty area", worst < 0.05,
          f"dark_band_rows={worst:.0%}")
    if multipage:
        check(f"{label} actually spans multiple pages", pages >= 2, f"pages={pages}")
        tops, lefts = [], []
        for pg in doc:
            t, l = _ink_bounds(pg)
            tops.append(t)
            lefts.append(l)
        # The real invariant for "spacing not disturbed across pages":
        #  (a) NO page is flush to the top edge (no sliced content / lost margin),
        #  (b) CONTINUATION pages (2..N) share the same top margin as each other,
        #  (c) page 1 may sit LOWER (a big name/header pushes its first ink down),
        #      but never HIGHER than the body margin,
        #  (d) the side (left) margin is identical on every page.
        cont = tops[1:]                       # continuation pages
        check(f"{label} no page flush to top edge", min(tops) >= MIN_TOP_MARGIN_MM,
              f"min_top={min(tops):.1f}mm tops={[round(x,1) for x in tops]}")
        check(f"{label} continuation pages share one top margin",
              (max(cont) - min(cont)) <= MAX_MARGIN_SPREAD_MM,
              f"cont_spread={max(cont)-min(cont):.1f}mm tops={[round(x,1) for x in tops]}")
        check(f"{label} page 1 not tighter than body margin",
              tops[0] >= min(cont) - 1.0,
              f"page1_top={tops[0]:.1f}mm body_top={min(cont):.1f}mm")
        # Side-margin check: continuation pages must share one left margin (a
        # real horizontal shift between body pages would be a bug). Page 1's
        # first ink is NOT compared, because its first element may be a
        # full-width banner sitting at the exact margin or an indented body
        # block — that's content layout, not the page margin (which @page fixes
        # identically for every page).
        contL = lefts[1:]
        check(f"{label} continuation pages share one left margin",
              (max(contL) - min(contL)) <= MAX_MARGIN_SPREAD_MM,
              f"cont_spread={max(contL)-min(contL):.1f}mm lefts={[round(x,1) for x in lefts]}")
    doc.close()


def test_docx(profile, content):
    label = f"[{profile}] DOCX"
    try:
        b = build_resume_docx(content)
    except Exception as e:
        check(f"{label} builds", False, str(e)[:120])
        return
    z = zipfile.ZipFile(io.BytesIO(b))
    check(f"{label} is valid OOXML (opens clean)", z.testzip() is None and "word/document.xml" in z.namelist())
    d = Document(io.BytesIO(b))
    txt = "\n".join(p.text for p in d.paragraphs)
    check(f"{label} has text", len(txt.strip()) > 200, f"len={len(txt.strip())}")
    rels = [r for r in d.part.rels.values() if "hyperlink" in r.reltype]
    check(f"{label} has real hyperlinks", len(rels) >= 1, f"links={len(rels)}")
    # custom link label (feature 1.4) appears as visible text
    has_label = any("mental-health-companion" in p.text for p in d.paragraphs) or \
        any("github" in (p.text.lower()) for p in d.paragraphs)
    check(f"{label} renders link label text", has_label)


def before_after_pagination():
    """Explicit before/after on the creative-bold multi-page spacing fix:
    strip the @media print block (the 'before' state) and show page 2 is flush,
    then with it (the 'after' state) show every page has a consistent margin."""
    print("\n-- BEFORE/AFTER: creative-bold multi-page spacing --")
    html_path = os.path.join(FIX, "creative-bold.threepage.html")
    if not os.path.exists(html_path):
        print("  (skipped: fixture missing)")
        return
    after_html = open(html_path, encoding="utf-8").read()
    # 'before' = remove the print pagination block we added
    before_html = re.sub(r"@media print\s*\{.*?\}\s*\}", "", after_html, flags=re.DOTALL)

    def page_tops(html):
        doc = fitz.open(stream=render_html_to_pdf(html), filetype="pdf")
        tops = [_ink_bounds(pg)[0] for pg in doc]
        doc.close()
        return tops

    b_tops = page_tops(before_html)
    a_tops = page_tops(after_html)
    print(f"  BEFORE page tops (mm): {[round(x,1) for x in b_tops]}  -> page 2 flush: {min(b_tops[1:]) < MIN_TOP_MARGIN_MM}")
    print(f"  AFTER  page tops (mm): {[round(x,1) for x in a_tops]}  -> all margined: {min(a_tops) >= MIN_TOP_MARGIN_MM}")
    check("creative-bold BEFORE reproduces the bug (page 2 flush)", min(b_tops[1:]) < MIN_TOP_MARGIN_MM)
    check("creative-bold AFTER fixes it (every page margined & consistent)",
          min(a_tops) >= MIN_TOP_MARGIN_MM and (max(a_tops) - min(a_tops)) <= MAX_MARGIN_SPREAD_MM)


def before_after_colored_bleed():
    """Before/after for the dark <html> background bleeding into the empty area
    below content (creative-bold). 'before' = drop the !important white override
    so the inline dark bg wins again."""
    print("\n-- BEFORE/AFTER: creative-bold colored-box bleed --")
    p = os.path.join(FIX, "creative-bold.onepage.html")
    if not os.path.exists(p):
        print("  (skipped)")
        return
    after = open(p, encoding="utf-8").read()
    before = after.replace("background: #fff !important", "background: #fff")

    def worst(html):
        doc = fitz.open(stream=render_html_to_pdf(html), filetype="pdf")
        v = max(_dark_bleed_band_fraction(pg) for pg in doc)
        doc.close()
        return v

    b, a = worst(before), worst(after)
    print(f"  BEFORE dark band in empty area: {b:.0%}  -> bug present: {b >= 0.05}")
    print(f"  AFTER  dark band in empty area: {a:.0%}  -> clean: {a < 0.05}")
    check("creative-bold BEFORE reproduces colored-box bleed", b >= 0.05)
    check("creative-bold AFTER removes the colored-box bleed", a < 0.05)


def main():
    htmls = sorted(glob.glob(os.path.join(FIX, "*.html")))
    if not htmls:
        print("No fixtures. Run: cd ../resume-ai-frontend && node scripts/_smoke-render.mjs")
        sys.exit(2)

    import json
    print("== PDF tests ==")
    for h in htmls:
        base = os.path.basename(h)[:-5]          # e.g. creative-bold.threepage
        slug, profile = base.rsplit(".", 1)
        html = open(h, encoding="utf-8").read()
        test_pdf(slug, profile, html, multipage=(profile == "threepage"))

    print("\n== DOCX tests ==")
    for pj in sorted(glob.glob(os.path.join(FIX, "*.input.json"))):
        profile = os.path.basename(pj).split(".")[0]
        content = json.load(open(pj, encoding="utf-8"))
        test_docx(profile, content)

    before_after_pagination()
    before_after_colored_bleed()

    total = len(PASS) + len(FAIL)
    print(f"\n==== RESULT: {len(PASS)}/{total} passed, {len(FAIL)} failed ====")
    if FAIL:
        print("FAILED:")
        for f in FAIL:
            print("  -", f)
        sys.exit(1)
    print("ALL TESTS PASSED")


if __name__ == "__main__":
    main()
